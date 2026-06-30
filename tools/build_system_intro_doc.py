from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "system_intro_assets"
DOCX_PATH = OUT_DIR / "ERP台账系统功能介绍.docx"


BLUE = "1F5FBF"
DARK = "0F172A"
INK = "1E293B"
MUTED = "64748B"
LIGHT_BG = "F6F8FB"
BORDER = "D9E2EF"
GREEN = "0F9F6E"
ORANGE = "D97706"
RED = "DC2626"


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def add_text(draw, xy, text, size=28, fill="#111827", bold=False, anchor=None):
    if isinstance(fill, str) and len(fill) == 6 and not fill.startswith("#"):
        fill = f"#{fill}"
    draw.text(xy, text, font=font(size, bold), fill=fill, anchor=anchor)


def rounded(draw, box, radius=18, fill="#FFFFFF", outline="#D9E2EF", width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#64748B", width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - 14 * direction, y2 - 8), (x2 - 14 * direction, y2 + 8)]
    else:
        direction = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 8, y2 - 14 * direction), (x2 + 8, y2 - 14 * direction)]
    draw.polygon(pts, fill=fill)


def make_architecture(path: Path):
    img = Image.new("RGB", (1600, 900), "#F8FAFC")
    d = ImageDraw.Draw(img)
    add_text(d, (80, 58), "ERP台账系统总体架构", 44, "#0F172A", True)
    add_text(d, (80, 112), "React + Vite 前端、FastAPI 服务、MySQL 台账数据库与 Excel 导入链路", 24, "#475569")

    # Layers
    layers = [
        (80, 180, 420, 680, "用户与前端层", "#E0F2FE"),
        (520, 180, 1080, 680, "后端服务层", "#EEF2FF"),
        (1180, 180, 1520, 680, "数据与运维层", "#ECFDF5"),
    ]
    for x1, y1, x2, y2, title, fill in layers:
        rounded(d, (x1, y1, x2, y2), 24, fill, "#CBD5E1", 2)
        add_text(d, (x1 + 28, y1 + 28), title, 30, "#0F172A", True)

    front_cards = [
        ("登录与角色权限", "管理员、业务、采购、销售、只读"),
        ("首页仪表盘", "KPI、趋势、部门排行、日志"),
        ("台账/订单/采购/销售", "查询、分页、录入、详情查看"),
    ]
    for i, (title, desc) in enumerate(front_cards):
        y = 260 + i * 120
        rounded(d, (120, y, 380, y + 82), 16, "#FFFFFF", "#BAE6FD")
        add_text(d, (145, y + 20), title, 24, "#075985", True)
        add_text(d, (145, y + 52), desc, 18, "#475569")

    api_cards = [
        ("认证与权限", "/api/auth、JWT、角色控制"),
        ("业务查询 API", "/api/ledgers、/orders、/purchases、/sales"),
        ("统计与视图", "项目汇总、订单财务汇总、首页指标"),
        ("导入与系统维护", "Excel解析、操作日志、备份记录"),
    ]
    for i, (title, desc) in enumerate(api_cards):
        y = 240 + i * 96
        rounded(d, (570, y, 1030, y + 70), 14, "#FFFFFF", "#C7D2FE")
        add_text(d, (596, y + 16), title, 23, "#3730A3", True)
        add_text(d, (596, y + 46), desc, 18, "#475569")

    data_cards = [
        ("MySQL 8.0", "项目、订单、采购、销售、财务明细"),
        ("Excel台账导入", "491条有效业务明细、原始行留痕"),
        ("审计与备份", "操作日志、导入批次、备份记录"),
    ]
    for i, (title, desc) in enumerate(data_cards):
        y = 260 + i * 120
        rounded(d, (1220, y, 1480, y + 82), 16, "#FFFFFF", "#BBF7D0")
        add_text(d, (1245, y + 20), title, 24, "#166534", True)
        add_text(d, (1245, y + 52), desc, 18, "#475569")

    arrow(d, (420, 430), (520, 430), "#2563EB", 5)
    arrow(d, (1080, 430), (1180, 430), "#059669", 5)
    add_text(d, (455, 392), "HTTP API", 18, "#2563EB", True)
    add_text(d, (1112, 392), "SQL/导入", 18, "#059669", True)
    img.save(path)


def make_workflow(path: Path):
    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    d = ImageDraw.Draw(img)
    add_text(d, (80, 58), "业务数据闭环", 44, "#0F172A", True)
    add_text(d, (80, 112), "从 Excel 初始化到订单、采购、销售、财务核对和报表汇总", 24, "#475569")

    steps = [
        ("Excel导入", "按批次读取台账\n保存原始行与校验结果", "#DBEAFE", "#1D4ED8"),
        ("订单建档", "项目编号、订单号、订单行\n支持一个项目多个订单", "#E0E7FF", "#4338CA"),
        ("采购跟踪", "供应商、采购合同\n收票、入库、多期付款", "#FEF3C7", "#B45309"),
        ("销售跟踪", "销售合同、开票\n多期回款、应收款", "#DCFCE7", "#047857"),
        ("财务核对", "应收、应付、毛利润\n关闭状态与异常预警", "#FFE4E6", "#BE123C"),
        ("首页报表", "KPI、趋势、部门排行\n最近操作日志", "#F1F5F9", "#334155"),
    ]
    positions = [(80, 270), (330, 270), (580, 270), (830, 270), (1080, 270), (1330, 270)]
    for idx, ((title, desc, fill, color), (x, y)) in enumerate(zip(steps, positions), start=1):
        rounded(d, (x, y, x + 190, y + 230), 24, fill, color, 3)
        d.ellipse((x + 24, y + 24, x + 72, y + 72), fill=color)
        add_text(d, (x + 48, y + 49), str(idx), 24, "#FFFFFF", True, anchor="mm")
        add_text(d, (x + 24, y + 92), title, 28, color, True)
        for line_no, line in enumerate(desc.split("\n")):
            add_text(d, (x + 24, y + 136 + line_no * 30), line, 20, "#334155")
        if idx < len(steps):
            arrow(d, (x + 195, y + 115), (x + 245, y + 115), "#94A3B8", 4)

    rounded(d, (180, 610, 1420, 835), 22, "#F8FAFC", "#CBD5E1", 2)
    add_text(d, (220, 666), "关键口径", 26, "#0F172A", True)
    notes = [
        "订单价值 = 数量 × 销售单价，优先保留台账原值",
        "应收款 = 订单价值 - 回款合计；应付账款 = 采购金额 - 已付款金额",
        "重复发生的开票、收票、付款、回款按 phase_no 建模，支持多期业务",
    ]
    for i, note in enumerate(notes):
        add_text(d, (220, 705 + i * 34), f"{i + 1}. {note}", 20, "#475569")
    img.save(path)


def make_ui_preview(path: Path):
    img = Image.new("RGB", (1600, 1000), "#F3F4F6")
    d = ImageDraw.Draw(img)
    # Sidebar
    d.rectangle((0, 0, 270, 1000), fill="#0F172A")
    rounded(d, (28, 30, 74, 76), 8, "#FFFFFF", "#FFFFFF", 1)
    add_text(d, (92, 42), "中通服供应链安徽", 18, "#FFFFFF", True)
    nav = ["首页仪表盘", "台账管理", "订单详情", "采购详情", "销售详情", "系统管理"]
    for i, item in enumerate(nav):
        y = 130 + i * 58
        fill = "#1D4ED8" if i == 0 else "#0F172A"
        d.rounded_rectangle((24, y, 246, y + 42), radius=8, fill=fill)
        add_text(d, (58, y + 11), item, 19, "#E2E8F0", True if i == 0 else False)

    # Header
    d.rectangle((270, 0, 1600, 70), fill="#FFFFFF")
    add_text(d, (310, 24), "系统核心 / 首页仪表盘", 21, "#2563EB", True)
    add_text(d, (1320, 24), "最后更新: 10:28:36", 18, "#64748B")

    # Content
    add_text(d, (300, 120), "仪表盘概览", 34, "#0F172A", True)
    add_text(d, (300, 164), "欢迎回来，这是今天的业务实时动态。", 20, "#64748B")
    metrics = [
        ("订单总金额", "¥1,286.4万", BLUE),
        ("毛利润", "¥214.8万", GREEN),
        ("订单总数", "491", BLUE),
        ("应收账款", "¥328.6万", RED),
        ("应付账款", "¥176.2万", ORANGE),
        ("已关闭订单", "128", GREEN),
    ]
    for i, (label, value, color) in enumerate(metrics):
        x = 300 + i * 205
        rounded(d, (x, 220, x + 180, 330), 16, "#FFFFFF", "#E2E8F0")
        add_text(d, (x + 20, 246), label, 18, "#64748B")
        add_text(d, (x + 20, 282), value, 28, color, True)

    rounded(d, (300, 370, 1040, 690), 18, "#FFFFFF", "#E2E8F0")
    add_text(d, (330, 400), "订单与毛利润趋势", 22, "#0F172A", True)
    for i in range(4):
        y = 470 + i * 45
        d.line((340, y, 1000, y), fill="#E2E8F0", width=2)
    pts = [(350, 620), (480, 560), (610, 590), (740, 500), (870, 460), (1000, 420)]
    d.line(pts, fill="#2563EB", width=5, joint="curve")
    for x, y in pts:
        d.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#2563EB")
    pts2 = [(350, 650), (480, 620), (610, 625), (740, 580), (870, 555), (1000, 520)]
    d.line(pts2, fill="#64748B", width=4)

    rounded(d, (1080, 370, 1520, 690), 18, "#FFFFFF", "#E2E8F0")
    add_text(d, (1110, 400), "部门订单排行", 22, "#0F172A", True)
    depts = [("市场一部", 0.92), ("市场二部", 0.73), ("政企事业部", 0.56), ("平台业务部", 0.42)]
    for i, (name, pct) in enumerate(depts):
        y = 465 + i * 54
        add_text(d, (1110, y), name, 19, "#334155")
        d.rounded_rectangle((1240, y + 8, 1490, y + 22), radius=7, fill="#E2E8F0")
        d.rounded_rectangle((1240, y + 8, 1240 + int(250 * pct), y + 22), radius=7, fill="#2563EB")

    rounded(d, (300, 730, 1520, 940), 18, "#FFFFFF", "#E2E8F0")
    add_text(d, (330, 762), "操作日志", 22, "#0F172A", True)
    headers = ["操作人", "操作模块", "详情", "状态", "操作时间"]
    xs = [330, 520, 720, 1160, 1320]
    for x, h in zip(xs, headers):
        add_text(d, (x, 812), h, 18, "#64748B", True)
    rows = [
        ("管理员", "导入任务", "完成 2026 市场部业务台账初始化", "成功", "2026-06-28"),
        ("业务人员", "订单详情", "录入订单交付信息", "成功", "2026-06-28"),
    ]
    for r, row in enumerate(rows):
        y = 852 + r * 38
        d.line((330, y - 12, 1490, y - 12), fill="#E2E8F0", width=1)
        for x, cell in zip(xs, row):
            add_text(d, (x, y), cell, 17, "#334155")
    img.save(path)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EF"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_style_font(style, size=None, color=None, bold=None, before=None, after=None, line_spacing=None):
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    if size:
        style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if before is not None:
        style.paragraph_format.space_before = Pt(before)
    if after is not None:
        style.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED)


def add_image(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F4F7")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_font(run, size=9.5, color=INK)
    doc.add_paragraph()
    return table


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    arch = ASSET_DIR / "architecture.png"
    flow = ASSET_DIR / "workflow.png"
    ui = ASSET_DIR / "ui_preview.png"
    make_architecture(arch)
    make_workflow(flow)
    make_ui_preview(ui)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style_font(doc.styles["Normal"], size=11, color=INK, before=0, after=6, line_spacing=1.10)
    set_style_font(doc.styles["Heading 1"], size=16, color=BLUE, bold=True, before=16, after=8)
    set_style_font(doc.styles["Heading 2"], size=13, color=BLUE, bold=True, before=12, after=6)
    set_style_font(doc.styles["Heading 3"], size=12, color="1F4D78", bold=True, before=8, after=4)
    set_style_font(doc.styles["List Bullet"], size=11, color=INK, before=0, after=4, line_spacing=1.167)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("ERP 台账系统功能介绍")
    set_font(run, size=24, color=DARK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("面向供应链公司业务台账的轻量 ERP 管理系统")
    set_font(r, size=12, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    meta.autofit = False
    meta_rows = [
        ("系统版本", "v1.0.0"),
        ("技术栈", "React + Vite 前端 / FastAPI 后端 / MySQL 8.0 数据库"),
        ("数据基础", "2026年市场部业务台账；主台账 87 列，491 条有效业务明细"),
    ]
    for row_idx, (label, value) in enumerate(meta_rows):
        for col_idx, text in enumerate((label, value)):
            cell = meta.rows[row_idx].cells[col_idx]
            cell.width = Inches(1.55 if col_idx == 0 else 4.95)
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(text)
            set_font(rr, size=10, color=INK, bold=(col_idx == 0))
    doc.add_paragraph()

    add_heading(doc, "1. 系统概述", 1)
    add_body(
        doc,
        "ERP 台账系统用于把原本分散在 Excel 台账中的订单、采购、交付、销售、开票、回款、付款等信息沉淀到数据库，并通过网页端提供查询、录入、汇总和维护能力。系统当前保持前端页面结构稳定，后端通过 API 向首页、台账、订单、采购、销售和系统管理模块提供真实数据。",
    )
    add_bullets(
        doc,
        [
            "统一业务主线：以项目编号、订单号、订单行作为核心业务粒度，支持一个项目编号对应多个订单号。",
            "覆盖财务闭环：把应收、应付、回款、付款、毛利润、关闭状态等指标统一计算和展示。",
            "保留导入留痕：Excel 导入后保存原始行、导入批次和校验结果，便于后续核对。",
            "面向本地部署：前端运行在 3000 端口，后端运行在 8000 端口，数据库使用 MySQL。",
        ],
    )
    add_image(doc, arch, "图 1  系统总体架构示意")

    add_heading(doc, "2. 用户角色与权限", 1)
    add_body(doc, "系统内置基于角色的权限控制。不同岗位可以进入同一套系统，但写入范围按职责拆分，减少误操作风险。")
    add_table(
        doc,
        ["角色", "主要职责", "典型权限"],
        [
            ("管理员", "系统维护、数据导入、用户管理、备份恢复", "全部模块；账号管理；系统导入与备份"),
            ("订单录入", "维护订单基础信息、订单行与批量导入", "订单详情录入；订单模板下载与导入"),
            ("采购录入", "维护采购合同、收票、入库、付款", "采购详情写入；采购合同与付款信息维护"),
            ("销售录入", "维护销售合同、开票、回款", "销售详情写入；销售合同、发票和回款维护"),
            ("只读查看", "查询台账、报表和系统数据", "读取首页、台账和业务明细"),
        ],
        [1.05, 2.45, 3.0],
    )

    add_heading(doc, "3. 核心功能模块", 1)
    add_heading(doc, "3.1 首页仪表盘", 2)
    add_body(
        doc,
        "首页提供经营状态的总览入口，展示订单总金额、毛利润、订单总数、应收账款、应付账款、已关闭订单数，并支持按部门筛选。趋势图和部门排行用于快速判断业务规模、利润变化和责任部门分布。",
    )
    add_image(doc, ui, "图 2  首页仪表盘界面示意")

    add_heading(doc, "3.2 台账管理", 2)
    add_body(
        doc,
        "台账管理是系统的总览查询页面，面向项目级别汇总。用户可以按项目编号、部门、客户经理、客户单位、订单号、订单状态和订单日期范围查询；列表展示订单金额、采购金额、回款合计、应收款和毛利润。点击查看后，可展开项目、订单、采购销售、收款付款等完整信息。",
    )

    add_heading(doc, "3.3 订单详情", 2)
    add_body(
        doc,
        "订单详情面向订单行维护，记录项目编号、订单号、订单日期、业务类型、客户单位、货物名称、规格型号、数量、单价、订单价值和交付数量。模块支持下载 CSV 模板、批量导入订单，以及按订单维护采购和交付信息。",
    )

    add_heading(doc, "3.4 采购详情", 2)
    add_body(
        doc,
        "采购详情围绕订单行跟踪供应商、采购合同、采购金额、收票和付款情况。后端支持按 order_line_id 查看采购明细，并分别新增采购合同、采购发票和多期付款记录。应付账款由采购金额和已付款金额汇总得到。",
    )

    add_heading(doc, "3.5 销售详情", 2)
    add_body(
        doc,
        "销售详情围绕销售合同、开票和回款管理，支持按项目、订单、客户经理和部门查询。明细接口返回合同、发票和回款记录，并支持多期回款。应收账款由订单价值与回款合计计算得到，便于财务核对和关闭状态判断。",
    )

    add_heading(doc, "3.6 系统管理", 2)
    add_body(
        doc,
        "系统管理模块包含账号与权限管理、操作日志、备份信息和系统刷新。管理员可创建账号并分配角色；操作日志记录导入、录入、备份等动作；备份信息用于查看和登记数据库备份文件。",
    )

    add_heading(doc, "4. 业务数据流", 1)
    add_body(
        doc,
        "系统以 Excel 台账作为初始化来源，导入后拆分为项目、订单、订单行、采购、交付、销售、发票、回款、付款等结构化数据，并通过汇总视图服务前端页面。重复发生的收票、付款、开票和回款使用 phase_no 表示期次，避免把多期业务硬编码成固定列。",
    )
    add_image(doc, flow, "图 3  业务数据闭环示意")

    add_heading(doc, "5. 主要接口与数据对象", 1)
    add_table(
        doc,
        ["模块", "主要接口/对象", "说明"],
        [
            ("健康检查", "/api/health", "返回数据库配置状态、导入行数和错误信息"),
            ("认证权限", "/api/auth/login、/api/auth/me、erp_user", "登录、当前用户、角色权限校验"),
            ("首页仪表盘", "/api/dashboard/*、v_project_ledger_summary", "汇总 KPI、趋势、部门排行和近期日志"),
            ("台账管理", "/api/ledgers、v_project_ledger_summary", "项目级汇总查询与分页"),
            ("订单详情", "/api/orders、order_line", "订单行查询、交付与金额字段展示"),
            ("采购详情", "/api/purchases、purchase_contract、purchase_invoice、purchase_payment", "采购合同、收票、多期付款和应付账款"),
            ("销售详情", "/api/sales、sales_contract、sales_invoice、sales_receipt", "销售合同、开票、多期回款和应收账款"),
            ("系统维护", "/api/logs、/api/backups、ledger_import_batch、ledger_raw_row", "操作日志、备份记录、导入批次和原始行留痕"),
        ],
        [1.15, 2.8, 2.55],
    )

    add_heading(doc, "6. 部署与运行方式", 1)
    add_body(doc, "系统采用前后端分离的本地部署方式。后端启动时会初始化表结构、创建默认管理员，并在业务表为空时尝试导入 Excel 台账。")
    add_table(
        doc,
        ["项目", "说明"],
        [
            ("前端", "React + Vite；默认访问地址 http://127.0.0.1:3000"),
            ("后端", "FastAPI + Uvicorn；默认接口地址 http://127.0.0.1:8000"),
            ("数据库", "MySQL 8.0；默认库名 erp_ledger，可通过 backend/.env 配置"),
            ("默认账号", "admin / admin123；正式使用前应修改 DEFAULT_ADMIN_PASSWORD 和 AUTH_SECRET"),
            ("启动命令", "npm run dev；npm run backend:dev"),
        ],
        [1.3, 5.2],
    )

    add_heading(doc, "7. 系统价值", 1)
    add_bullets(
        doc,
        [
            "把 Excel 静态台账升级为可查询、可录入、可审计的业务系统。",
            "把订单、采购、销售、财务口径串成一条闭环，减少人工汇总和重复核对。",
            "保留原始导入数据，便于和财务、业务部门按来源行追溯差异。",
            "通过角色权限控制不同岗位的写入边界，降低业务数据维护风险。",
            "保留现有前端结构，后续可以继续增量扩展审批、导出、报表和主数据模块。",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("ERP 台账系统功能介绍")
    set_font(fr, size=9, color=MUTED)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
